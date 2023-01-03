from goose3 import Goose
from parsers import parse
import glob, os, textwrap
from datetime import date
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Note: Filepaths must end with "/"
DOWNLOADS = "/Users/username/Downloads/"
OUTPUT_FOLDER = "/path/to/output/folder/"
NUM_FILES = 4


def main():
    all_files = glob.glob(DOWNLOADS + "*.html")
    relevant_files = sorted(all_files, key=os.path.getctime)[-NUM_FILES:]
    articles = []
    print("*" * 80)
    for i, file in enumerate(relevant_files):
        html = open(file).read()
        title, domain = extract_metadata(html)
        pars, removed_list = parse(domain, html)
        print(f"Article {i+1}/{NUM_FILES}")
        if pars is None:
            print(f'Could not parse article: "{title}" from {domain}')
            print("Press any key to continue.")
            input()
            print("*" * 80)
            continue
        for r in removed_list:
            print(f"    REMOVED: {r}")
        print("----")
        print(f"Title:    {title}")
        print(f"Site:     {domain}")
        wordcount = sum([len(x) for x in pars])
        print(f"# words:  {wordcount}")
        def indented_print(s):
            return "\n    ".join(textwrap.wrap(s))
        print(f"1st Par:\n    {indented_print(pars[0][:200])}...")
        print(f"Last Par:\n    ...{indented_print(pars[-1][-200:])}")
        print("Remove last paragraph? (y/N):")

        while input().lower() in {"y", "yes"}:
            pars = pars[:-1]
            print(f"Last Par:\n    ...{indented_print(pars[-1][-200:])}")
            print("Remove last paragraph? (y/N):")
        
        article = "\n\n".join(pars)
        articles.append(f"{title}\n{domain}\n\n{article}")
        print("#"*80)
    write_doc(articles)

def write_doc(articles):
    doc = Document()
    for a in articles[:-1]:
        doc.add_paragraph(a)
        doc.add_page_break()
    doc.add_paragraph(articles[-1])

    doc.save(f"{OUTPUT_FOLDER}/{date.today()}-output.docx")

def extract_metadata(html_input):
    # Returns title and site, because that is easy
    # Also, I forget the actual metadata required
    g = Goose()
    goose_article = g.extract(raw_html=html_input)
    return [goose_article.title, goose_article.domain]

# # Currently not useful (should extract text...)
# def extract_hyperlinks(filename):
#     document = Document(filename)
#     rels = document.part.rels

#     names = []
#     urls = []
#     for r in rels:
#         if rels[r].reltype == RT.HYPERLINK:
#             names.append(rels[r].__dict__)
#             urls.append(rels[r]._target)
#     return(names[0])

#     # def iter_hyperlink_rels(rels):
#     #     for rel in rels:
#     #         if rels[rel].reltype == RT.HYPERLINK:
#     #             yield rels[rel]._target      

#     # print(iter_hyperlink_rels(rels))

if __name__ == "__main__":
    main()